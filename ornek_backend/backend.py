import os
import logging
import time
from io import BytesIO
from fastapi import FastAPI, File, UploadFile, Query, HTTPException
from fastapi.responses import JSONResponse
import uvicorn
from PyPDF2 import PdfReader
import docx

# LangChain bileşenleri:
# - Document: Dosya içeriklerinin doküman nesnelerine dönüştürülmesi.
# - ChatOpenAI: LLM çağrılarını yönetir.
# - ChatPromptTemplate, SystemMessagePromptTemplate, HumanMessagePromptTemplate: Prompt oluşturma araçları.
# - RecursiveCharacterTextSplitter: Büyük metinleri parçalamak için kullanılır.
# - HuggingFaceEmbeddings: Metinleri vektörleştirmek için HuggingFace tabanlı embedding modeli.
# - Chroma: Vektör veritabanı, benzerlik araması için kullanılır.
from langchain.docstore.document import Document
from langchain_community.chat_models import ChatOpenAI  
from langchain.prompts.chat import (
    ChatPromptTemplate,
    SystemMessagePromptTemplate,
    HumanMessagePromptTemplate
)
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_huggingface import HuggingFaceEmbeddings  
from langchain_community.vectorstores import Chroma

# LM Studio ve model ayarları
LM_STUDIO_ENDPOINT = "http://192.168.88.100:1234/v1"
MODEL_IDENTIFIER = "llama-3.2-3b-instruct" # modelleri çeşitlendirelim
LLM_TOKEN_LIMIT = 4096  # LLM'in kabul edebileceği maksimum token sayısı

# Logging yapılandırması: Uygulama genelinde hata ve bilgi mesajlarını loglamak için kullanılır.
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("stlc_manager")

app = FastAPI()

def read_file_content(file_path: str) -> str:
    """
    Belirtilen dosya yolundan metin içeriğini okur.
    Kullanım amacı: system_message.txt gibi yapılandırma dosyalarını okumak.
    """
    try:
        with open(file_path, "r", encoding="utf-8") as file:
            return file.read().strip()
    except FileNotFoundError:
        logger.error(f"File not found: {file_path}")
        return None

def pdf_to_text(file_stream: BytesIO) -> str:
    """
    PDF dosyasındaki metni çıkarır.
    Kullanım amacı: Yüklenen PDF dosyalarından metin elde etmek.
    """
    text = ""
    try:
        reader = PdfReader(file_stream)
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n"
    except Exception as e:
        logger.error("Error reading PDF: %s", e)
    return text

def docx_to_text(file_stream: BytesIO) -> str:
    """
    DOCX dosyasındaki metni çıkarır.
    Kullanım amacı: Yüklenen DOCX dosyalarından metin elde etmek.
    """
    try:
        doc = docx.Document(file_stream)
        return "\n".join([para.text for para in doc.paragraphs])
    except Exception as e:
        logger.error("Error reading DOCX: %s", e)
        return ""

def txt_to_text(file_stream: BytesIO) -> str:
    """
    TXT dosyasındaki metni çıkarır.
    Kullanım amacı: Yüklenen TXT dosyalarından metin elde etmek.
    """
    try:
        return file_stream.read().decode('utf-8')
    except Exception as e:
        logger.error("Error reading TXT: %s", e)
        return ""

def extract_text(upload_file: UploadFile) -> str:
    """
    Yüklenen dosyanın uzantısına göre uygun metin çıkarma fonksiyonunu çağırır.
    Kullanım amacı: Farklı dosya formatları (PDF, DOCX, TXT) için ortak metin çıkarımı.
    """
    ext = os.path.splitext(upload_file.filename)[1].lower()
    content = upload_file.file.read()
    upload_file.file.seek(0)
    if ext == ".pdf":
        return pdf_to_text(BytesIO(content))
    elif ext == ".docx":
        return docx_to_text(BytesIO(content))
    elif ext == ".txt":
        return txt_to_text(BytesIO(content))
    else:
        logger.warning(f"Unsupported file type: {ext}")
        return ""

def sanitize_text(text: str) -> str:
    """
    Giriş metninde gereksiz boşlukları ve karakterleri temizler.
    Kullanım amacı: Metin girişlerinin temiz ve standart formata getirilmesi.
    """
    return " ".join(text.split())

def count_tokens(text: str) -> int:
    """
    Basit token sayımı yapar (kelimeler üzerinden).
    Kullanım amacı: LLM token limiti kontrolü için metindeki token sayısını belirlemek.
    """
    return len(text.split())

def determine_chunk_size(text: str, base_size=1000, min_size=500):
    """
    Dinamik olarak metnin token sayısına göre chunk boyutunu ayarlar.
    Eğer metin, LLM token limitini aşıyorsa, chunk boyutunu düşürerek LLM'e uygun parçalara böler.
    Kullanım amacı: LLM'in token limitine uyum sağlamak.
    """
    tokens = count_tokens(text)
    if tokens > LLM_TOKEN_LIMIT:
        factor = tokens / LLM_TOKEN_LIMIT
        new_size = max(int(base_size / factor), min_size)
        logger.info(f"Adjusting chunk size: base {base_size} -> {new_size} (token count: {tokens})")
        return new_size
    return base_size

def retry_invoke(chain, inputs, retries=3, timeout=10):
    """
    LLM çağrısı sırasında hata alınırsa, belirli sayıda yeniden deneme yapar.
    Kullanım amacı: Ağ veya LLM kaynaklı geçici hatalarda sistemin kararlı çalışmasını sağlamak.
    """
    last_exception = None
    for attempt in range(retries):
        try:
            # chain.invoke fonksiyonunun timeout desteklemesi durumunda timeout parametresi de entegre edilebilir.
            return chain.invoke(inputs)
        except Exception as e:
            logger.error(f"LLM invocation failed on attempt {attempt+1}: {e}")
            last_exception = e
            time.sleep(1)  # Yeniden deneme öncesi kısa bekleme
    raise last_exception

@app.get("/models")
def list_models():
    """
    Sistemde kullanılabilir modellerin listesini döner.
    Kullanım amacı: UI veya diğer servislerin, mevcut model seçeneklerini öğrenmesi.
    """
    return JSONResponse(content={"available_models": [MODEL_IDENTIFIER]})

@app.post("/upload")
async def upload_documents(
    files: list[UploadFile] = File(...),
    model_name: str = Query(..., description="Kullanılacak model. Örneğin: 'llama-3.2-3b-instruct'")
):
    """
    Yüklenen dosyalardan metin çıkarımı yapar, gerekli sanitizasyonu ve chunking işlemlerini gerçekleştirir,
    ardından LLM'e gönderilecek promptu oluşturup, sonuçları döner.
    
    Ek Açıklamalar:
    - Dosya formatına göre metin çıkarımı yapılır.
    - Çıkarılan metin sanitize edilerek temizlenir.
    - Metin, LLM'in token limitine uygun parçalara (chunk) ayrılır; dinamik chunk boyutu ayarlanır.
    - Embedding modeli ile dokümanlar vektörleştirilir ve benzerlik araması yapılır.
    - LLM çağrısı yapılmadan önce prompt token limiti kontrol edilir; gerekirse metin kısaltılır.
    - LLM çağrısı için retry mekanizması kullanılır.
    - Eğer yüklenen dökümanların toplam token sayısı LLM'in işleyebileceği limiti aşıyorsa, bu durum uyarı olarak loglanır ve yanıt içerisinde bildirilir.
    """
    if model_name != MODEL_IDENTIFIER:
        logger.error("Invalid model name provided: %s", model_name)
        raise HTTPException(status_code=400, detail=f"Geçersiz model adı. Mevcut model: {MODEL_IDENTIFIER}")
    
    aggregated_text = ""
    # Her dosyadan metin çıkarımı, sanitizasyon ve birleşik metin oluşturma
    for file in files:
        text = extract_text(file)
        text = sanitize_text(text)
        aggregated_text += text + "\n\n"
    
    if not aggregated_text.strip():
        logger.error("No valid text extracted from files.")
        raise HTTPException(status_code=400, detail="Yüklenen dosyalardan geçerli metin çıkarılamadı.")
    
    # Uyarı: Toplam token sayısı LLM_TOKEN_LIMIT'i aşıyorsa uyarı ver.
    total_tokens = count_tokens(aggregated_text)
    warning_message = ""
    if total_tokens > LLM_TOKEN_LIMIT:
        warning_message = f"Uploaded documents' token count ({total_tokens}) exceeds the processing limit ({LLM_TOKEN_LIMIT}). Some content may be truncated."
        logger.warning(warning_message)
    
    # System prompt: LLM'e genel davranış ve stil verecek yönergeler okunuyor.
    system_message = read_file_content("system_message.txt")
    if system_message is None:
        raise HTTPException(status_code=400, detail="system_message.txt dosyası bulunamadı.")
    
    # Dinamik chunk boyutu ayarı: Metnin token sayısına göre chunk boyutu belirlenir.
    chunk_size = determine_chunk_size(aggregated_text, base_size=1000, min_size=500)
    text_splitter = RecursiveCharacterTextSplitter(chunk_size=chunk_size, chunk_overlap=100)
    chunks = text_splitter.split_text(aggregated_text)
    docs = [Document(page_content=chunk) for chunk in chunks]
    
    # Embedding model kullanımı:
    # HuggingFace tabanlı embedding modeli ile dokümanlar vektörleştirilir.
    # Bu sayede, belirli bir sorguya göre (örneğin test planı oluşturma) en uygun metin parçaları seçilebilir.
    embedding_model = HuggingFaceEmbeddings(model_name="BAAI/bge-small-en")
    vectorstore = Chroma.from_documents(docs, embedding_model, collection_name="uploaded_docs")
    
    # Sorgu ifadesi: Belirli bir test planı oluşturma isteğini temsil eder.
    query_str = "Create a detailed test plan based on the documents"
    relevant_docs = vectorstore.similarity_search(query_str, k=3)
    retrieved_text = "\n\n".join([doc.page_content for doc in relevant_docs])
    
    # LLM Token Limit Kontrolü:
    # system_message ve retrieved_text birleşimi LLM'e gönderilmeden önce token sayısı kontrol edilir.
    # Eğer token sayısı limiti aşıyorsa, retrieved_text kısaltılarak limitle uyum sağlanır.
    combined_prompt = f"{system_message}\n\n{retrieved_text}\n\nBased on these documents, create a detailed test plan."
    token_count = count_tokens(combined_prompt)
    if token_count > LLM_TOKEN_LIMIT:
        logger.warning(f"Combined prompt token count ({token_count}) exceeds limit ({LLM_TOKEN_LIMIT}). Truncating retrieved text.")
        allowed_tokens = LLM_TOKEN_LIMIT - count_tokens(system_message) - 50  # 50 token buffer ekleniyor
        retrieved_tokens = retrieved_text.split()[:allowed_tokens]
        retrieved_text = " ".join(retrieved_tokens)
    
    # Prompt oluşturma: System ve Human mesajları, LLM'e gönderilecek promptu oluşturur.
    prompt_template = ChatPromptTemplate.from_messages([
        SystemMessagePromptTemplate.from_template("{system_message}"),
        HumanMessagePromptTemplate.from_template("{retrieved_text}\n\nBased on these documents, create a detailed test plan.")
    ])
    
    # LLM çağrısı için ChatOpenAI kullanımı:
    # Bu, LLM çağrısını yapar ve modelin yanıtını alır.
    llm = ChatOpenAI(
        model_name=MODEL_IDENTIFIER,
        openai_api_base=LM_STUDIO_ENDPOINT,
        openai_api_key="not-needed",
        temperature=0.7
    )
    
    # LLM çağrısı zinciri: prompt_template ve LLM birleşimi.
    chain = prompt_template | llm
    try:
        # retry_invoke, LLM çağrısı sırasında hata alınırsa yeniden deneme mekanizması sağlar.
        response = retry_invoke(chain, {"system_message": system_message, "retrieved_text": retrieved_text}, retries=3, timeout=10)
    except Exception as e:
        logger.error("LLM invocation failed after retries: %s", e)
        raise HTTPException(status_code=500, detail="LLM çağrısı sırasında hata meydana geldi.")
    
    # Yanıtı string formatına çevirme
    result = response.content if hasattr(response, "content") else str(response)
    
    # API yanıtında sonucu ve varsa uyarı mesajını döndür.
    return JSONResponse(content={"result": result, "warning": warning_message})

if __name__ == "__main__":
    uvicorn.run("backend:app", host="0.0.0.0", port=8000, reload=True)
